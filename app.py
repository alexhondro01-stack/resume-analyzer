import streamlit as st
import os
import tempfile
import re
import requests
import json
from pathlib import Path
from datetime import datetime
from bs4 import BeautifulSoup
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from agents import ResumeCrew
from fpdf import FPDF
from tools import extract_text
from dotenv import load_dotenv

# --- CONFIGURATION ---
BASE_DIR = Path(__file__).resolve().parent
ENV_PATH = BASE_DIR / ".env"
if ENV_PATH.exists():
    load_dotenv(dotenv_path=ENV_PATH, override=True) 
else:
    load_dotenv(override=True)

st.set_page_config(page_title="AI Resume Matcher", layout="wide", page_icon="🚀")

# Initialize session state
if 'step' not in st.session_state: st.session_state.step = 1
if 'provider' not in st.session_state: st.session_state.provider = "OpenAI"
if 'jd_text' not in st.session_state: st.session_state.jd_text = ""
if 'analysis_result' not in st.session_state: st.session_state.analysis_result = None
if 'optimized_result' not in st.session_state: st.session_state.optimized_result = None
if 'selected_improvements' not in st.session_state: st.session_state.selected_improvements = []

# --- HELPERS ---
def get_api_key(provider_name):
    env_var = "OPENAI_API_KEY" if provider_name == "OpenAI" else "GOOGLE_API_KEY"
    key = os.environ.get(env_var, os.getenv(env_var, ""))
    return key.strip().strip("'").strip('"')

def extract_score(text):
    m = re.search(r"(\d{1,3})\s*/\s*100", str(text))
    if not m: 
        m = re.search(r"(?:Score|Match|Likelihood):\s*(\d{1,3})", str(text), re.IGNORECASE)
    return int(m.group(1)) if m else 0

def clean_for_latin1(text):
    if not text: return ""
    chars = {
        '\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'", 
        '\u201c': '"', '\u201d': '"', '\u2022': '*', '\u2026': '...',
        '\u00a0': ' ', '\ufb01': 'fi', '\ufb02': 'fl'
    }
    for k, v in chars.items(): text = text.replace(k, v)
    return text.encode('latin-1', 'ignore').decode('latin-1')

def create_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    safe_text = clean_for_latin1(text)
    clean_text = safe_text.split("TRANSFORMATION LOG")[0].replace("REVISED RESUME", "").strip()
    lines = clean_text.split('\n')
    
    # Standard Professional Header
    pdf.set_font("Helvetica", 'B', 16)
    if lines:
        # Title/Name in standard black
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 10, lines[0].strip(), ln=True, align='L')
        pdf.ln(5)
    
    pdf.set_font("Helvetica", size=11)
    for line in lines[1:]:
        if not line.strip(): 
            pdf.ln(2)
            continue
        # Section Headers (Bold, slight spacing)
        if len(line) < 50 and (line.isupper() or line.endswith(':')):
            pdf.ln(4)
            pdf.set_font("Helvetica", 'B', 12)
            pdf.cell(0, 10, line.strip(), ln=True)
            pdf.set_font("Helvetica", size=11)
        else:
            pdf.multi_cell(0, 6, line.strip())
            
    return pdf.output(dest='S').encode('latin-1')

def create_docx(text):
    doc = Document()
    clean_text = text.split("TRANSFORMATION LOG")[0].replace("REVISED RESUME", "").strip()
    lines = clean_text.split('\n')
    
    # Standard Formatting
    if lines:
        h = doc.add_heading(lines[0].strip(), level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0) # Black

    for line in lines[1:]:
        if line.strip():
            # Heuristic for Section Headers
            if len(line) < 50 and (line.isupper() or line.endswith(':')):
                p = doc.add_heading(line.strip(), level=1)
                run = p.runs[0]
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                p = doc.add_paragraph(line.strip())
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def fetch_jd(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36'}
        res = requests.get(url, headers=headers, timeout=12)
        res.raise_for_status()
        soup = BeautifulSoup(res.text, 'html.parser')
        for element in soup(["script", "style", "nav", "footer", "header", "aside", "form", "button"]):
            element.extract()
        content_selectors = ["article", "main", ".job-description", ".show-more-less-html__markup", "#jobDescriptionText", "[class*='jobDescription']", ".description__text", ".careers-job-description"]
        main_content = None
        for selector in content_selectors:
            found = soup.select_one(selector)
            if found:
                main_content = found
                break
        target = main_content if main_content else soup.body
        relevant_text = []
        priority_keywords = ["responsibility", "requirement", "qualification", "skill", "experience", "bonus", "stack", "technology", "about the role"]
        noise_keywords = ["equal opportunity", "inclusive", "cookie", "copyright", "all rights reserved", "privacy policy", "terms of service", "follow us"]
        if target:
            for block in target.find_all(['p', 'li', 'h1', 'h2', 'h3', 'h4', 'div']):
                text = block.get_text(strip=True)
                if len(text) < 15: continue 
                if any(noise in text.lower() for noise in noise_keywords): continue
                is_priority = any(pw in text.lower() for pw in priority_keywords)
                if is_priority or block.name == 'li' or len(text) > 40:
                    clean_block = re.sub(r'\s+', ' ', text)
                    relevant_text.append(clean_block)
        final_lines = []
        for line in relevant_text:
            if not final_lines or line[:30] != final_lines[-1][:30]:
                final_lines.append(line)
        return "\n\n".join(final_lines)[:8000]
    except Exception as e:
        return f"Scraping Error: {str(e)}. Please copy/paste the JD text manually for best results."

def parse_feedback_advanced(text):
    data = {
        "matches": [],
        "requirements": [],
        "qualifications": []
    }
    current_section = None
    lines = str(text).split('\n')
    for line in lines:
        l = line.lower()
        if "section: matches" in l: current_section = "matches"
        elif "section: job requirements" in l: current_section = "requirements"
        elif "section: qualification gaps" in l: current_section = "qualifications"
        if line.strip().startswith(('-', '*', '•', '1.')):
            clean_line = line.strip().lstrip('- *•1.2.3. ')
            if current_section == "matches":
                data["matches"].append(clean_line)
            elif current_section in ["requirements", "qualifications"]:
                gap_type = "REPURPOSE"
                if "[missing]" in l: gap_type = "MISSING"
                elif "[repurpose]" in l: gap_type = "REPURPOSE"
                priority = "low"
                if "[high]" in l: priority = "high"
                elif "[medium]" in l: priority = "medium"
                elif "[low]" in l: priority = "low"
                display_text = re.sub(r'\[.*?\]', '', clean_line).strip()
                if len(display_text) > 2:
                    data[current_section].append({
                        "text": display_text,
                        "type": gap_type,
                        "priority": priority,
                        "raw": clean_line
                    })
    return data

# --- SIDEBAR ---
with st.sidebar:
    st.header("⚙️ Settings")
    st.session_state.provider = st.selectbox("LLM Provider", ["OpenAI", "Gemini"], index=0 if st.session_state.provider == "OpenAI" else 1)
    o_key, g_key = get_api_key("OpenAI"), get_api_key("Gemini")
    st.caption(f"OpenAI: {'✅' if o_key else '❌'} | Gemini: {'✅' if g_key else '❌'}")
    with st.expander("🔑 Manual Key Overwrite"):
        key_input = st.text_input("Enter API Key", type="password")
        if st.button("Apply Key"):
            os.environ["OPENAI_API_KEY" if st.session_state.provider == "OpenAI" else "GOOGLE_API_KEY"] = key_input.strip()
            st.rerun()
    st.divider()
    if st.button("🗑️ Reset All"):
        st.session_state.clear()
        st.rerun()

# --- STEPPER ---
s1, s2, s3 = st.columns(3)
with s1: st.markdown(f"**{'🔵' if st.session_state.step==1 else '✅'} 1. Setup**")
with s2: st.markdown(f"**{'🔵' if st.session_state.step==2 else '✅' if st.session_state.step>2 else '⚪'} 2. Analysis**")
with s3: st.markdown(f"**{'🔵' if st.session_state.step==3 else '⚪'} 3. Results**")
st.divider()

# --- STEP 1: SETUP ---
if st.session_state.step == 1:
    st.subheader("📄 Your Resume")
    up = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])
    if up: st.session_state.original_filename = up.name

    st.divider()
    st.subheader("💼 Job Posting :orange[(BETA)]")
    url = st.text_input("Job URL (LinkedIn, Indeed, etc.)")
    if st.button("Fetch Relevant Content (Beta)") and url:
        with st.spinner("Analyzing web content..."): st.session_state.jd_text = fetch_jd(url)
    st.session_state.jd_text = st.text_area("Job Description Content", value=st.session_state.jd_text, height=250)

    st.divider()
    if st.button("Generate Match Report →", type="primary", use_container_width=True):
        current_key = get_api_key(st.session_state.provider)
        if up and len(st.session_state.jd_text) > 50 and current_key:
            with st.spinner("Analyzing your profile..."):
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{up.name.split('.')[-1]}") as t:
                    t.write(up.getvalue()); st.session_state.resume_path = t.name
                st.session_state.original_text = extract_text(st.session_state.resume_path)
                crew = ResumeCrew(st.session_state.resume_path, st.session_state.jd_text, st.session_state.provider)
                st.session_state.analysis_result = str(crew.analyze())
                st.session_state.step = 2; st.rerun()
        else: st.error("Please provide Resume, JD, and a valid API Key.")

# --- STEP 2: ANALYSIS & INPUT ---
elif st.session_state.step == 2:
    score = extract_score(st.session_state.analysis_result)
    
    # Dashboard Header
    sc_col, info_col = st.columns([1, 2])
    color = "#ef4444" if score < 50 else "#f97316" if score < 75 else "#22c55e"
    with sc_col:
        st.markdown(f"<div style='background:{color}; color:white; padding:30px; border-radius:15px; text-align:center;'><p style='margin:0; font-size:14px;'>MATCH SCORE</p><h1 style='margin:0; font-size:64px;'>{score}%</h1></div>", unsafe_allow_html=True)
    with info_col:
        st.markdown("### 🎯 Optimization Roadmap")
        st.write("Review the matches and gaps below. For **missing** items, you must provide context to include them.")
        st.progress(score/100)

    st.divider()
    
    # Advanced Parsing
    data = parse_feedback_advanced(st.session_state.analysis_result)
    
    # 1. Matches
    st.markdown("#### ✅ Matched Strengths")
    if data["matches"]:
        cols = st.columns(2)
        for i, m in enumerate(data["matches"]):
            cols[i%2].success(m)
    else: st.info("No direct matches found.")
    
    st.divider()

    # 2. Gaps & Inputs
    st.markdown("#### 🛠️ Gaps & Improvements")
    
    selected_items = []
    
    def render_gap_list(items, section_name):
        if not items: return
        st.subheader(section_name)
        priority_order = {"high": 0, "medium": 1, "low": 2}
        items.sort(key=lambda x: priority_order.get(x['priority'], 3))
        
        for i, gap in enumerate(items):
            p_color = "red" if gap['priority'] == "high" else "orange" if gap['priority'] == "medium" else "blue"
            type_icon = "❓ MISSING" if gap['type'] == "MISSING" else "🔄 REPURPOSE"
            
            with st.container():
                col_check, col_details = st.columns([0.05, 0.95])
                k_base = f"{section_name}_{i}"
                with col_details:
                    st.markdown(f":{p_color}[**[{gap['priority'].upper()}]**] **{gap['text']}** ({type_icon})")
                    if gap['type'] == "MISSING":
                        user_input = st.text_input(f"Describe your experience with {gap['text']} to include it:", key=f"in_{k_base}", placeholder="e.g. Used for 2 years in project X...")
                        if user_input:
                            selected_items.append(f"Add missing skill '{gap['text']}' with context: {user_input}")
                            st.caption("✅ Included for optimization")
                        else: st.caption("⚠️ Input required to select")
                    else:
                        if st.checkbox(f"Optimize phrasing for {gap['text']}", value=True, key=f"cb_{k_base}"):
                            selected_items.append(f"Repurpose existing experience to highlight '{gap['text']}'")

    render_gap_list(data["requirements"], "Job Requirements")
    render_gap_list(data["qualifications"], "Qualifications")

    st.divider()
    b1, b2 = st.columns([1, 3])
    with b1: 
        if st.button("← Go Back"): st.session_state.step = 1; st.rerun()
    with b2:
        if st.button("Apply & Rewrite Resume ✨", type="primary", use_container_width=True):
            if not selected_items:
                st.warning("Please select at least one item (and provide input for missing skills).")
            else:
                st.session_state.selected_improvements = selected_items
                with st.spinner("Rewriting resume..."):
                    crew = ResumeCrew(st.session_state.resume_path, st.session_state.jd_text, st.session_state.provider)
                    st.session_state.optimized_result = str(crew.optimize(", ".join(selected_items)))
                    st.session_state.step = 3; st.rerun()

# --- STEP 3: RESULTS ---
elif st.session_state.step == 3:
    st.subheader("✨ Final Optimized Resume")
    
    full_text = st.session_state.optimized_result
    resume_part, log_part = full_text.split("TRANSFORMATION LOG") if "TRANSFORMATION LOG" in full_text else (full_text, "")
    
    st.text_area("Resume Content (Editable)", value=resume_part.replace("REVISED RESUME", "").strip(), height=600, key="final_edit")
    with st.expander("View AI Changes Log"): st.write(log_part.strip())
    
    with st.popover("📥 Download", use_container_width=True):
        docx_data = create_docx(st.session_state.optimized_result)
        st.download_button("Word (.docx)", docx_data, file_name="Optimized.docx", use_container_width=True)
        try:
            pdf_data = create_pdf(st.session_state.optimized_result)
            st.download_button("PDF (.pdf)", pdf_data, file_name="Optimized.pdf", use_container_width=True)
        except: st.error("PDF Error")
        
    if st.button("Start New Analysis"): st.session_state.step = 1; st.rerun()